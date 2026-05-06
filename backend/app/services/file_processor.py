import hashlib
from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader
from app.config import settings


def get_file_hash(file_content: bytes) -> str:
    """Generate MD5 hash of file content to prevent duplicates."""
    return hashlib.md5(file_content).hexdigest()


def validate_file(filename: str, file_size: int) -> bool:
    """Validate file extension and size."""
    allowed_types = settings.allowed_file_types.split(",")
    extension = Path(filename).suffix.lstrip(".").lower()

    if extension not in allowed_types:
        return False

    max_bytes = settings.max_file_size_mb * 1024 * 1024
    if file_size > max_bytes:
        return False

    return True


def parse_docx(file_content: bytes) -> str:
    """Extract all text from DOCX file."""
    from io import BytesIO
    doc = DocxDocument(BytesIO(file_content))

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text_parts.append(" | ".join(row_text))

    return "\n".join(text_parts)


def parse_xlsx(file_content: bytes) -> Dict[str, Any]:
    """Extract structured data from XLSX file."""
    from io import BytesIO
    workbook = load_workbook(BytesIO(file_content))

    result = {
        "sheets": {},
        "raw_text": []
    }

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        rows = []

        for row in sheet.iter_rows(values_only=True):
            rows.append(row)

        result["sheets"][sheet_name] = rows

        for row in rows:
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                result["raw_text"].append(row_text)

    return result


def parse_pdf(file_content: bytes) -> str:
    """Extract all text from PDF file."""
    from io import BytesIO
    reader = PdfReader(BytesIO(file_content))

    text_parts = []
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text()
        if text.strip():
            text_parts.append(f"[Page {page_num + 1}]\n{text}")

    return "\n".join(text_parts)


def process_file(file_content: bytes, filename: str) -> tuple[str, str, str]:
    """
    Process uploaded file and extract text content.
    Returns: (document_type, extracted_text, file_hash)
    """
    extension = Path(filename).suffix.lstrip(".").lower()
    file_hash = get_file_hash(file_content)

    if extension == "docx":
        text = parse_docx(file_content)
        return ("property_notes", text, file_hash)

    elif extension == "xlsx":
        data = parse_xlsx(file_content)
        text = "\n".join(data["raw_text"])
        return ("revenue_data", text, file_hash)

    elif extension == "pdf":
        text = parse_pdf(file_content)
        return ("document", text, file_hash)

    else:
        raise ValueError(f"Unsupported file type: {extension}")
